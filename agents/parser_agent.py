from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path
from typing import Any, List, Optional

from docx import Document
from pydantic import BaseModel, ConfigDict, Field

from backend.client import YANDEX_CLOUD_FOLDER, YANDEX_CLOUD_MODEL, client


class Goal(BaseModel):
    """Цель проекта."""

    goal: str = Field(description="Текст цели")
    confidence: float = Field(ge=0.0, le=1.0, description="Уверенность (0-1)")


class Result(BaseModel):
    """Результат проекта."""

    result: str = Field(description="Текст результата")
    confidence: float = Field(ge=0.0, le=1.0, description="Уверенность (0-1)")
    measurable: bool = Field(description="Измеримый результат")


class Timeline(BaseModel):
    """Сроки проекта."""

    start_date: Optional[date] = Field(None, description="Дата начала")
    end_date: Optional[date] = Field(None, description="Дата окончания")
    duration: Optional[str] = Field(None, description="Длительность")
    confidence: float = Field(ge=0.0, le=1.0, description="Уверенность")


class ProjectData(BaseModel):
    """Основная модель с целями, сроками и результатами."""

    model_config = ConfigDict(
        json_schema_extra={
            "example": {
                "project_name": "Разработка системы кредитного скоринга",
                "goals": [
                    {
                        "goal": "Разработать модель кредитного скоринга",
                        "confidence": 0.95,
                    }
                ],
                "results": [
                    {
                        "result": "Обученная модель с ROC-AUC >= 0.75",
                        "confidence": 0.9,
                        "measurable": True,
                    }
                ],
                "timeline": {
                    "start_date": "2026-03-01",
                    "end_date": "2026-05-31",
                    "confidence": 0.95,
                },
            }
        }
    )

    project_name: str = Field(description="Название проекта")
    goals: List[Goal] = Field(default_factory=list, description="Цели проекта")
    results: List[Result] = Field(default_factory=list, description="Результаты проекта")
    timeline: Optional[Timeline] = Field(None, description="Сроки проекта")


class ProjectParser:
    """Агент для извлечения структуры проекта из DOCX-паспорта."""

    def __init__(
        self,
        *,
        max_document_chars: int = 15000,
        temperature: float = 0.1,
    ) -> None:
        if not YANDEX_CLOUD_FOLDER or not YANDEX_CLOUD_MODEL:
            raise ValueError(
                "Не заданы YANDEX_CLOUD_FOLDER или YANDEX_CLOUD_MODEL в окружении."
            )

        self.model = f"gpt://{YANDEX_CLOUD_FOLDER}/{YANDEX_CLOUD_MODEL}"
        self.max_document_chars = max_document_chars
        self.temperature = temperature

    def read_docx(self, file_path: str | Path) -> str:
        """Читает текст и таблицы из DOCX файла."""
        path = Path(file_path).expanduser()
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {path}")

        doc = Document(str(path))
        text: list[str] = []

        for paragraph in doc.paragraphs:
            value = paragraph.text.strip()
            if value:
                text.append(value)

        for table in doc.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " | ")
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    text.append(" || ".join(cells))

        return "\n".join(text)

    def parse(self, file_path: str | Path) -> ProjectData:
        """Парсит DOCX и возвращает валидированные проектные данные."""
        text = self.read_docx(file_path)
        if not text:
            raise ValueError("Не удалось извлечь текст из файла")

        response_text = self._ask_llm(text[: self.max_document_chars])
        result_dict = self._parse_json(response_text)
        return ProjectData.model_validate(result_dict)

    def _ask_llm(self, document_text: str) -> str:
        schema = json.dumps(ProjectData.model_json_schema(), ensure_ascii=False, indent=2)
        prompt = f"""
Проанализируй паспорт проекта и извлеки из него структурированные данные.

Текст документа:
{document_text}

Верни только JSON-объект, который соответствует этой JSON Schema:
{schema}

Правила извлечения:
1. project_name бери из названия документа или полей "Название проекта", "Название", "Карточка инициативы".
2. goals бери из разделов "Цель проекта", "Целевое состояние", "Главная цель".
3. results бери из разделов "Ожидаемый результат", "Планируемые результаты", "Что должно быть получено" и критериев приемки, если они описывают конечные проверяемые результаты.
4. timeline бери из разделов или полей "Плановый срок", "Сроки", "Срок реализации", "Период выполнения". Даты возвращай в формате YYYY-MM-DD.
5. confidence оценивает явность источника:
   - 0.9-1.0: прямое явное указание;
   - 0.7-0.89: информация следует из текста;
   - 0.5-0.69: предположение.
6. measurable показывает, можно ли проверить результат численно или по четкому артефакту.
7. Если информации нет, верни пустой массив для goals/results и null для timeline.
"""

        response = client.chat.completions.create(
            model=self.model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Ты эксперт по извлечению данных из проектной документации. "
                        "Отвечай только валидным JSON без markdown."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            temperature=self.temperature,
            response_format={"type": "json_object"},
        )

        return response.choices[0].message.content or "{}"

    @staticmethod
    def _parse_json(response_text: str) -> dict[str, Any]:
        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            match = re.search(r"\{.*\}", response_text, flags=re.DOTALL)
            if not match:
                raise
            return json.loads(match.group(0))
