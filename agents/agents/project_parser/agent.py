from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from agents.domain.project_document import ProjectData
from agents.infrastructure.llm import get_llm_adapter

from .prompts import PARSER_SYSTEM_PROMPT, build_parser_prompt


class ProjectParser:
    """Агент для извлечения структуры проекта из DOCX-паспорта."""

    def __init__(
        self,
        *,
        max_document_chars: int = 15000,
        temperature: float = 0.1,
    ) -> None:
        self.llm = get_llm_adapter()
        self.max_document_chars = max_document_chars
        self.temperature = temperature

    def read_docx(self, file_path: str | Path) -> str:
        """Читает текст и таблицы из DOCX файла."""
        path = Path(file_path).expanduser()
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {path}")

        doc = Document(str(path))
        text: list[str] = []

        for paragraph in doc.paragraphs:
            value = paragraph.text.strip()
            if value:
                text.append(value)

        for table in doc.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " | ")
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    text.append(" || ".join(cells))

        return "\n".join(text)

    async def parse(self, file_path: str | Path) -> ProjectData:
        """Парсит DOCX и возвращает валидированные проектные данные."""
        text = await asyncio.to_thread(self.read_docx, file_path)
        if not text:
            raise ValueError("Не удалось извлечь текст из файла")

        return await self._ask_llm(text[: self.max_document_chars])

    async def _ask_llm(self, document_text: str) -> ProjectData:
        return await self.llm.parse_pydantic(
            response_model=ProjectData,
            system_prompt=PARSER_SYSTEM_PROMPT,
            user_prompt=build_parser_prompt(document_text),
            temperature=self.temperature,
        )
