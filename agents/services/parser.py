from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from agents.domain.project_document import ProjectData
from agents.infrastructure.llm import get_llm_adapter


class ProjectParser:
    """Агент для извлечения структуры проекта из DOCX-паспорта."""

    def __init__(
        self,
        *,
        max_document_chars: int = 15000,
        temperature: float = 0.1,
    ) -> None:
        self.llm = get_llm_adapter()
        self.max_document_chars = max_document_chars
        self.temperature = temperature

    def read_docx(self, file_path: str | Path) -> str:
        """Читает текст и таблицы из DOCX файла."""
        path = Path(file_path).expanduser()
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {path}")

        doc = Document(str(path))
        text: list[str] = []

        for paragraph in doc.paragraphs:
            value = paragraph.text.strip()
            if value:
                text.append(value)

        for table in doc.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " | ")
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    text.append(" || ".join(cells))

        return "\n".join(text)

    async def parse(self, file_path: str | Path) -> ProjectData:
        """Парсит DOCX и возвращает валидированные проектные данные."""
        text = await asyncio.to_thread(self.read_docx, file_path)
        if not text:
            raise ValueError("Не удалось извлечь текст из файла")

        return await self._ask_llm(text[: self.max_document_chars])

    async def _ask_llm(self, document_text: str) -> ProjectData:
        prompt = f"""
Проанализируй паспорт проекта и извлеки из него структурированные данные.

Текст документа:
{document_text}

Верни ответ по строгому Pydantic-контракту ProjectData.

Правила извлечения:
1. project_name бери из названия документа или полей "Название проекта", "Название", "Карточка инициативы".
2. goals бери из разделов "Цель проекта", "Целевое состояние", "Главная цель".
3. results бери из разделов "Ожидаемый результат", "Планируемые результаты", "Что должно быть получено" и критериев приемки, если они описывают конечные проверяемые результаты.
4. timeline бери из разделов или полей "Плановый срок", "Сроки", "Срок реализации", "Период выполнения". Даты возвращай в формате YYYY-MM-DD.
5. confidence оценивает явность источника:
   - 0.9-1.0: прямое явное указание;
   - 0.7-0.89: информация следует из текста;
   - 0.5-0.69: предположение.
6. measurable показывает, можно ли проверить результат численно или по четкому артефакту.
7. Если информации нет, верни пустой массив для goals/results и null для timeline.
"""

        return await self.llm.parse_pydantic(
            response_model=ProjectData,
            system_prompt=(
                "Ты эксперт по извлечению данных из проектной документации. "
                "Отвечай только структурой ProjectData."
            ),
            user_prompt=prompt,
            temperature=self.temperature,
        )
